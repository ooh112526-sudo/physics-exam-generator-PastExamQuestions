import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import requests

# 題型對照表
TYPE_MAP_ZH_TO_EN = {"單選": "Single", "多選": "Multi", "填充": "Fill", "題組": "Group"}
TYPE_MAP_EN_TO_ZH = {v: k for k, v in TYPE_MAP_ZH_TO_EN.items()}

def get_image_bytes(q):
    """嘗試取得題目的圖片 (Base64 或 URL)"""
    if q.image_data: return q.image_data
    if q.image_url:
        try:
            response = requests.get(q.image_url, timeout=3)
            if response.status_code == 200: return response.content
        except: return None
    return None

def set_font(doc, style_name='Normal'):
    """設定中文字型"""
    style = doc.styles[style_name]
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

def generate_word_files(selected_questions, config=None):
    """
    生成 Word 檔案
    config: {
        'title': str,
        'teacher_version': bool, (教用卷: 含答案解析)
        'student_version': bool, (學用卷: 僅題目)
        'answer_version': bool   (答案卷: 僅簡答表)
    }
    """
    if config is None:
        config = {'title': '物理科試題卷', 'teacher_version': True, 'student_version': True, 'answer_version': True}

    outputs = {}
    
    # 1. 準備文件物件
    doc_teacher = docx.Document() if config.get('teacher_version') else None
    doc_student = docx.Document() if config.get('student_version') else None
    doc_answer = docx.Document() if config.get('answer_version') else None

    # 設定字型 & 標題
    title = config.get('title', '物理科試題卷')
    for doc in [doc_teacher, doc_student, doc_answer]:
        if doc:
            set_font(doc)
            # 標題
            h = doc.add_paragraph(title)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.runs[0].font.size = Pt(16)
            h.runs[0].bold = True
            doc.add_paragraph(f"範圍：綜合練習  總題數：{len(selected_questions)}題")
            doc.add_paragraph("-" * 80)

    # 重新編號計數器
    q_counter = 1
    
    # 用於答案卷的暫存 [(題號, 答案)]
    ans_list = []

    for q in selected_questions:
        # 計算此題佔用的題號範圍 (如果是題組，可能佔多個)
        current_range_str = str(q_counter)
        if q.is_group_parent:
            end_num = q_counter + len(q.sub_questions) - 1
            current_range_str = f"{q_counter}-{end_num}"
        
        # --- 處理教用卷與學用卷 ---
        for doc, is_teacher in [(doc_teacher, True), (doc_student, False)]:
            if not doc: continue
            
            p = doc.add_paragraph()
            
            # 題頭: [111-中模...] 【單選】 1. 
            # 僅教用卷顯示 exam_code 方便老師溯源
            code_str = f"[{q.exam_code}] " if (is_teacher and q.exam_code) else ""
            type_str = f"【{TYPE_MAP_EN_TO_ZH.get(q.type, q.type)}】"
            
            prefix = ""
            if q.is_group_parent:
                prefix = f"{current_range_str}題為題組 {code_str}\n"
            else:
                prefix = f"{q_counter}. {code_str}{type_str} "
            
            runner = p.add_run(prefix + q.content.strip())
            if is_teacher: runner.bold = False # 教用卷不一定要粗體，視閱讀舒適度
            
            # 圖片
            img_bytes = get_image_bytes(q)
            if img_bytes:
                try:
                    doc.add_paragraph().add_run().add_picture(io.BytesIO(img_bytes), width=Inches(3.0))
                except: pass
            
            # 選項
            if q.type in ['Single', 'Multi'] and q.options:
                opts = q.options
                # 簡單排版：長度短的併排，長的換行
                max_len = max([len(str(o)) for o in opts]) if opts else 0
                if max_len < 15 and len(opts) > 0:
                    doc.add_paragraph("　".join(opts))
                else:
                    for opt in opts: doc.add_paragraph(f"{opt}")
            
            # 題組的子題處理
            if q.is_group_parent:
                local_counter = q_counter
                for sub in q.sub_questions:
                    # 子題文字
                    sp = doc.add_paragraph()
                    sub_type = f"【{TYPE_MAP_EN_TO_ZH.get(sub.type, sub.type)}】"
                    sp.add_run(f"{local_counter}. {sub_type} {sub.content.strip()}")
                    
                    # 子題選項
                    if sub.type in ['Single', 'Multi'] and sub.options:
                        s_opts = sub.options
                        s_max_len = max([len(str(o)) for o in s_opts]) if s_opts else 0
                        if s_max_len < 15 and len(s_opts) > 0:
                            doc.add_paragraph("　".join(s_opts))
                        else:
                            for opt in s_opts: doc.add_paragraph(f"{opt}")
                    
                    # 答案與解析 (僅教用卷)
                    if is_teacher:
                        ans_p = doc.add_paragraph()
                        ans_p.add_run(f"答：{sub.answer}").bold = True
                        if sub.solution:
                            sol_p = doc.add_paragraph()
                            sol_p.add_run("解：").bold = True
                            sol_p.add_run(sub.solution)
                            
                    # 記錄到答案卷
                    ans_list.append((local_counter, sub.answer))
                    local_counter += 1
                doc.add_paragraph("-" * 20) # 題組分隔線
                
            else:
                # 一般題目的答案與解析 (僅教用卷)
                if is_teacher:
                    ans_p = doc.add_paragraph()
                    ans_p.add_run(f"答：{q.answer}").bold = True
                    if q.solution:
                        sol_p = doc.add_paragraph()
                        sol_p.add_run("解：").bold = True
                        sol_p.add_run(q.solution)
                elif q.type == 'Fill':
                    doc.add_paragraph("答：____________")

                # 記錄到答案卷
                ans_list.append((q_counter, q.answer))
            
            doc.add_paragraph("") # 空行分隔

        # 更新計數器
        if q.is_group_parent:
            q_counter += len(q.sub_questions)
        else:
            q_counter += 1

    # --- 處理答案卷 ---
    if doc_answer:
        # 表格排版，每行 5 題
        table = doc_answer.add_table(rows=1, cols=10)
        table.style = 'Table Grid'
        
        # 填入 Header
        row = table.rows[0].cells
        for i in range(5):
            row[i*2].text = "題號"
            row[i*2+1].text = "答案"
            
        # 填入數據
        current_row_idx = 0
        col_idx = 0
        for num, ans in ans_list:
            if col_idx >= 5: # 換行
                table.add_row()
                current_row_idx += 1
                col_idx = 0
            
            cells = table.rows[current_row_idx].cells
            cells[col_idx*2].text = str(num)
            cells[col_idx*2+1].text = str(ans)
            col_idx += 1

    # 輸出 ByteIO
    if doc_teacher:
        bio = io.BytesIO()
        doc_teacher.save(bio)
        bio.seek(0)
        outputs['teacher'] = bio
        
    if doc_student:
        bio = io.BytesIO()
        doc_student.save(bio)
        bio.seek(0)
        outputs['student'] = bio
        
    if doc_answer:
        bio = io.BytesIO()
        doc_answer.save(bio)
        bio.seek(0)
        outputs['answer'] = bio
        
    return outputs
