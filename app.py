import streamlit as st
import docx
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
import random
import io
import pandas as pd
import time
import base64
import requests 
from PIL import Image

# 安全載入 streamlit_cropper
try:
    from streamlit_cropper import st_cropper 
except ImportError:
    st_cropper = None 
except Exception:
    st_cropper = None

import os
import datetime
import uuid
import json
from google.cloud import firestore
from google.cloud import storage
import google.auth 
from google.oauth2 import service_account

import smart_importer

st.set_page_config(page_title="物理題庫系統 (Pro)", layout="wide", page_icon="🧲")

# [新功能] 題型對照表 (中文 <-> 英文)
TYPE_MAP_ZH_TO_EN = {"單選": "Single", "多選": "Multi", "填充": "Fill", "題組": "Group"}
TYPE_MAP_EN_TO_ZH = {v: k for k, v in TYPE_MAP_ZH_TO_EN.items()}
TYPE_OPTIONS = ["單選", "多選", "填充", "題組"]

# ==========================================
# 雲端資料庫與儲存模組 (CloudManager) - 保持不變
# ==========================================
class CloudManager:
    # ... (此處 CloudManager 程式碼保持不變，直接使用前一版的內容即可) ...
    # 為了節省篇幅，這裡省略 CloudManager 的完整定義，請確保複製完整版本
    def __init__(self):
        self.bucket_name = os.getenv("GCS_BUCKET_NAME", "physics-exam-assets")
        self.db = None
        self.storage_client = None
        self.has_connection = False
        self.connection_error = ""
        self.project_id = None
        self.credentials = None 

        try:
            service_account_json = os.getenv("GCP_SERVICE_ACCOUNT_JSON")
            if service_account_json:
                try:
                    service_account_json = service_account_json.strip()
                    if service_account_json.startswith("'") and service_account_json.endswith("'"):
                         service_account_json = service_account_json[1:-1]
                    service_account_info = json.loads(service_account_json)
                    self.credentials = service_account.Credentials.from_service_account_info(service_account_info)
                    self.project_id = service_account_info.get("project_id")
                    if not self.project_id: self.project_id = os.getenv("GCP_PROJECT_ID")
                    self.db = firestore.Client(credentials=self.credentials, project=self.project_id)
                    self.storage_client = storage.Client(credentials=self.credentials, project=self.project_id)
                    self.has_connection = True
                    if self.has_connection: self._ensure_bucket_exists()
                    return
                except Exception as e: print(f"JSON Env Error: {e}")

            try:
                if "gcp_service_account" in st.secrets:
                    service_account_info = st.secrets["gcp_service_account"]
                    self.credentials = service_account.Credentials.from_service_account_info(service_account_info)
                    self.project_id = service_account_info.get("project_id")
                    self.db = firestore.Client(credentials=self.credentials, project=self.project_id)
                    self.storage_client = storage.Client(credentials=self.credentials, project=self.project_id)
                    self.has_connection = True
                    return 
            except: pass

            self.project_id = (os.getenv("GCP_PROJECT_ID") or os.getenv("GOOGLE_CLOUD_PROJECT"))
            if not self.project_id:
                try: self.credentials, project_id_from_auth = google.auth.default(); self.project_id = project_id_from_auth
                except: pass

            if self.project_id:
                if self.credentials:
                    self.db = firestore.Client(credentials=self.credentials, project=self.project_id)
                    self.storage_client = storage.Client(credentials=self.credentials, project=self.project_id)
                else:
                    self.db = firestore.Client(project=self.project_id)
                    self.storage_client = storage.Client(project=self.project_id)
                self.has_connection = True
            else:
                try: self.db = firestore.Client(); self.storage_client = storage.Client(); self.has_connection = True
                except: pass
            
            if self.has_connection: self._ensure_bucket_exists()

        except Exception as e:
            self.connection_error = str(e)
            print(f"Cloud Init Error: {e}")

    def _ensure_bucket_exists(self):
        if not self.storage_client: return
        try:
            bucket = self.storage_client.bucket(self.bucket_name)
            if not bucket.exists(): bucket.create(location="us-central1") 
        except: pass

    def get_storage_usage(self):
        if not self.storage_client: return 0
        try:
            bucket = self.storage_client.bucket(self.bucket_name)
            return sum(blob.size for blob in bucket.list_blobs() if blob.size)
        except: return 0

    def upload_bytes(self, file_bytes, filename, folder="uploads", content_type=None):
        if not self.storage_client: return None
        try:
            bucket = self.storage_client.bucket(self.bucket_name)
            unique_name = f"{folder}/{int(datetime.datetime.now().timestamp())}_{str(uuid.uuid4())[:8]}_{filename}"
            blob = bucket.blob(unique_name)
            blob.upload_from_string(file_bytes, content_type=content_type)
            
            url = blob.public_url
            try:
                if self.credentials and hasattr(self.credentials, 'service_account_email'):
                     url = blob.generate_signed_url(version="v4", expiration=datetime.timedelta(days=7), method="GET", service_account_email=self.credentials.service_account_email, access_token=self.credentials.token)
                else:
                    url = blob.generate_signed_url(version="v4", expiration=datetime.timedelta(days=7), method="GET")
            except: pass
            return url, unique_name
        except: return None, None

    def download_blob(self, blob_name):
        if not self.storage_client or not blob_name: return None
        try:
            bucket = self.storage_client.bucket(self.bucket_name)
            blob = bucket.blob(blob_name)
            return blob.download_as_bytes()
        except: return None

    # --- File/Question Management ---
    def check_file_exists(self, filename):
        if not self.db: return None
        docs = self.db.collection("exam_files").where("filename", "==", filename).limit(1).stream()
        for doc in docs: d = doc.to_dict(); d['id'] = doc.id; return d
        return None

    def save_file_record(self, file_info, overwrite_id=None):
        if not self.db: return False
        doc_id = overwrite_id if overwrite_id else str(uuid.uuid4())
        file_info["id"] = doc_id; file_info["updated_at"] = datetime.datetime.now()
        self.db.collection("exam_files").document(doc_id).set(file_info)
        return True

    def load_file_records(self):
        if not self.db: return []
        files = []
        docs = self.db.collection("exam_files").order_by("updated_at", direction=firestore.Query.DESCENDING).stream()
        for doc in docs: files.append(doc.to_dict())
        return files

    def delete_file_record(self, file_id):
        if self.db: self.db.collection("exam_files").document(file_id).delete()

    def update_file_status(self, file_id, status):
        if self.db: self.db.collection("exam_files").document(file_id).update({"ai_status": status})

    def save_question(self, question_dict):
        if not self.db: return False
        if question_dict.get("image_data_b64"):
            try:
                img_bytes = base64.b64decode(question_dict["image_data_b64"])
                fname = f"q_{question_dict.get('id')}.png"
                img_url, _ = self.upload_bytes(img_bytes, fname, folder="question_images", content_type="image/png")
                if img_url: question_dict["image_url"] = img_url; del question_dict["image_data_b64"]
            except: pass
        self.db.collection("questions").document(question_dict["id"]).set(question_dict)
        return True

    def load_questions(self):
        if not self.db: return []
        qs = []
        docs = self.db.collection("questions").order_by("id").stream()
        for doc in docs: qs.append(doc.to_dict())
        return qs

    def delete_question(self, doc_id):
        if self.db: self.db.collection("questions").document(doc_id).delete()

cloud_manager = CloudManager()

# ... (Question 類別保持不變) ...
class Question:
    def __init__(self, q_type, content, options=None, answer=None, original_id=0, image_data=None, 
                 source="一般試題", chapter="未分類", unit="", db_id=None, 
                 parent_id=None, is_group_parent=False, sub_questions=None, image_url=None,
                 source_file_id=None):
        self.id = db_id if db_id else str(int(time.time()*1000)) + str(random.randint(0, 999))
        self.type = q_type 
        self.source = source
        self.chapter = chapter
        self.unit = unit
        self.content = content
        self.options = options if options else []
        self.answer = answer
        self.image_data = image_data 
        self.image_url = image_url   
        self.parent_id = parent_id 
        self.is_group_parent = is_group_parent 
        self.sub_questions = sub_questions if sub_questions else [] 
        self.source_file_id = source_file_id

    def to_dict(self):
        img_str = None
        if self.image_data: img_str = base64.b64encode(self.image_data).decode('utf-8')
        subs = [q.to_dict() for q in self.sub_questions] if self.sub_questions else []
        return {
            "id": self.id, "type": self.type, "source": self.source, "chapter": self.chapter,
            "content": self.content, "options": self.options, "answer": self.answer,
            "image_data_b64": img_str, "image_url": self.image_url,
            "parent_id": self.parent_id, "is_group_parent": self.is_group_parent,
            "sub_questions": subs, "source_file_id": self.source_file_id
        }

    @staticmethod
    def from_dict(data):
        img_bytes = None
        if data.get("image_data_b64"):
            try: img_bytes = base64.b64decode(data["image_data_b64"])
            except: pass
        q = Question(
            q_type=data.get("type", "Single"), content=data.get("content", ""),
            options=data.get("options", []), answer=data.get("answer", ""),
            original_id=0, image_data=img_bytes, image_url=data.get("image_url"),
            source=data.get("source", ""), chapter=data.get("chapter", "未分類"),
            db_id=data.get("id"), parent_id=data.get("parent_id"),
            is_group_parent=data.get("is_group_parent", False),
            source_file_id=data.get("source_file_id")
        )
        if data.get("sub_questions"):
            q.sub_questions = [Question.from_dict(sub) for sub in data["sub_questions"]]
        return q

if 'question_pool' not in st.session_state:
    st.session_state['question_pool'] = []
    try:
        data = cloud_manager.load_questions()
        if data: st.session_state['question_pool'] = [Question.from_dict(d) for d in data]
    except: pass

if 'file_queue' not in st.session_state: st.session_state['file_queue'] = {}

# ... (Utility Functions 保持不變) ...
def get_image_bytes(q):
    if q.image_data: return q.image_data
    if q.image_url:
        try:
            response = requests.get(q.image_url, timeout=3)
            if response.status_code == 200: return response.content
        except: return None
    return None

def generate_word_files(selected_questions):
    exam_doc = docx.Document()
    ans_doc = docx.Document()
    style = exam_doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')
    
    exam_doc.add_heading('物理科 試題卷', 0)
    ans_doc.add_heading('物理科 答案卷', 0)
    
    q_counter = 1
    
    def write_single_question(doc, q, idx_str):
        p = doc.add_paragraph()
        # [新功能] 題型顯示中文
        type_badge_zh = TYPE_MAP_EN_TO_ZH.get(q.type, q.type)
        type_label = f"【{type_badge_zh}】"
        src_label = f"[{q.source}] " if q.source and not q.parent_id else "" 
        
        runner = p.add_run(f"{idx_str}. {src_label}{type_label} {q.content.strip()}")
        runner.bold = True
        
        img_bytes = get_image_bytes(q)
        if img_bytes:
            try:
                img_p = doc.add_paragraph()
                run = img_p.add_run()
                run.add_picture(io.BytesIO(img_bytes), width=Inches(2.5))
            except: pass

        if q.type in ['Single', 'Multi'] and q.options:
            opts = q.options
            max_len = max([len(str(o)) for o in opts]) if opts else 0
            if max_len < 10 and len(opts) > 0:
                doc.add_paragraph("　　".join(opts))
            elif max_len < 25 and len(opts) > 0 and len(opts) % 2 == 0:
                table = doc.add_table(rows=(len(opts) // 2), cols=2)
                table.autofit = True
                for i, opt in enumerate(opts):
                    table.cell(i // 2, i % 2).text = opt
                doc.add_paragraph("")
            else:
                for opt in opts:
                    doc.add_paragraph(f"{opt}")
        elif q.type == 'Fill':
            doc.add_paragraph("答：______________________")
        doc.add_paragraph("") 

    for q in selected_questions:
        if q.is_group_parent:
            write_single_question(exam_doc, q, f"{q_counter}-{q_counter + len(q.sub_questions) - 1} 為題組")
            for sub_q in q.sub_questions:
                write_single_question(exam_doc, sub_q, str(q_counter))
                ans_p = ans_doc.add_paragraph()
                ans_p.add_run(f"{q_counter}. {sub_q.answer}")
                q_counter += 1
        else:
            write_single_question(exam_doc, q, str(q_counter))
            ans_p = ans_doc.add_paragraph()
            ans_p.add_run(f"{q_counter}. {q.answer}")
            q_counter += 1
        
    exam_io = io.BytesIO()
    ans_io = io.BytesIO()
    exam_doc.save(exam_io)
    ans_doc.save(ans_io)
    exam_io.seek(0)
    ans_io.seek(0)
    return exam_io, ans_io

def process_single_file(filename, api_key, file_id_in_db=None):
    if filename not in st.session_state['file_queue']: return
    info = st.session_state['file_queue'][filename]
    info['status'] = 'processing'
    
    with st.spinner(f"正在分析 {filename}... (AI 思考中，請稍候)"):
        # 修復：優先使用 blob_name 下載，解決過期問題
        file_bytes = info.get('data')
        blob_name = info.get('blob_name')
        
        # 如果記憶體中沒有，嘗試從 Cloud Storage 下載
        if not file_bytes and blob_name:
            file_bytes = cloud_manager.download_blob(blob_name)
            if file_bytes:
                st.session_state['file_queue'][filename]['data'] = file_bytes
        
        if not file_bytes:
            st.error("無法讀取檔案內容，請確認檔案是否已上傳。")
            info['status'] = 'error'
            return

        res = smart_importer.parse_with_gemini(file_bytes, info['type'], api_key)
    
    if isinstance(res, dict) and "error" in res:
        info['status'] = 'error'
        info['error_msg'] = res['error']
        st.error(f"{filename} 辨識失敗: {res['error']}")
    else:
        info['status'] = 'done'
        info['result'] = res
        if file_id_in_db:
            cloud_manager.update_file_status(file_id_in_db, "已辨識")
        
        st.success(f"{filename} 辨識完成！")
        st.session_state['just_processed_file'] = filename
        st.info("💡 請切換至「📝 AI匯入校對」分頁開始編輯")
        
    st.rerun()

# ==========================================
# Interface
# ==========================================
st.title("🧲 物理題庫系統 Pro (Cloud Storage)")

with st.sidebar:
    st.header("設定")
    env_api_key = os.getenv("GOOGLE_API_KEY", "")
    api_key_input = st.text_input("Gemini API Key", value=env_api_key, type="password", key="sidebar_api_key")
    
    if cloud_manager.has_connection:
        st.success("☁️ Cloud: 已連線")
        if cloud_manager.bucket_name:
            st.caption(f"Bucket: {cloud_manager.bucket_name}")
    else:
        st.warning(f"☁️ Cloud: 未連線")
        if cloud_manager.connection_error:
            st.caption(f"錯誤: {cloud_manager.connection_error}")
            if "No secrets found" in cloud_manager.connection_error:
                st.info("Secrets 未設定，請改用環境變數 GCP_SERVICE_ACCOUNT_JSON")

    st.divider()
    st.metric("題庫總數", len(st.session_state['question_pool']))
    
    if cloud_manager.has_connection:
        st.divider()
        try:
            total_bytes = cloud_manager.get_storage_usage()
            total_mb = total_bytes / (1024 * 1024)
            limit_mb = 1024.0 # 1GB
            percentage = min(total_mb / limit_mb, 1.0)
            
            st.write("📊 **雲端儲存空間**")
            st.progress(percentage)
            st.caption(f"已使用: {total_mb:.2f} MB / 1 GB")
            
            if percentage > 0.9:
                st.warning("⚠️ 容量即將額滿！")
        except:
            st.caption("無法取得容量資訊")

    if st.button("強制儲存至雲端", key="sidebar_force_save"):
        if cloud_manager.has_connection:
            progress_bar = st.progress(0)
            total = len(st.session_state['question_pool'])
            for i, q in enumerate(st.session_state['question_pool']):
                cloud_manager.save_question(q.to_dict())
                progress_bar.progress((i + 1) / total)
            st.success("儲存完成！")

# Tabs
tab_upload_process, tab_files, tab_review, tab_bank = st.tabs(["🧠 考古題上傳", "📂 檔案管理及AI辨識", "📝 AI匯入校對", "📚 題庫管理與試卷輸出"])

# === Tab 1: 考古題上傳 ===
with tab_upload_process:
    st.markdown("### 📤 上傳新考古題")
    st.info("請先選擇檔案，設定各自的標籤後，系統將自動重新命名並上傳。")
    
    uploaded_files = st.file_uploader("支援 .pdf, .docx", type=['pdf', 'docx'], accept_multiple_files=True)
    
    if uploaded_files:
        st.divider()
        st.subheader("設定檔案資訊")
        
        if 'upload_configs' not in st.session_state:
            st.session_state['upload_configs'] = {}

        with st.expander("批次設定 (一次套用給下方所有檔案)"):
            c_batch1, c_batch2, c_batch3, c_batch4 = st.columns(4)
            with c_batch1: b_type = st.selectbox("統一類型", ["學測", "分科", "北模", "中模", "全模", "其他"], key="batch_type")
            with c_batch2: b_year = st.text_input("統一年度", value="112", key="batch_year")
            with c_batch3: b_exam_no = st.selectbox("統一考試次別", ["第一次", "第二次", "第三次", "正式考試"], key="batch_no")
            with c_batch4: 
                if st.button("全部套用"):
                    for uf in uploaded_files:
                        st.session_state['upload_configs'][uf.name] = {
                            "type": b_type,
                            "year": b_year,
                            "exam_no": b_exam_no
                        }
                    st.success("已套用！")

        files_to_upload = []
        for i, f in enumerate(uploaded_files):
            current_config = st.session_state['upload_configs'].get(f.name, {
                "type": "學測", "year": "112", "exam_no": "正式考試"
            })
            
            with st.container():
                c1, c2, c3, c4 = st.columns([3, 2, 2, 2])
                with c1: 
                    st.markdown(f"**{i+1}. {f.name}**")
                    ext = f.name.split('.')[-1]
                    new_name = f"{current_config['year']}-{current_config['type']}-{current_config['exam_no']}.{ext}"
                    st.caption(f"➝ `{new_name}`")
                
                with c2: 
                    new_type = st.selectbox("類型", ["學測", "分科", "北模", "中模", "全模", "其他"], 
                                          index=["學測", "分科", "北模", "中模", "全模", "其他"].index(current_config['type']),
                                          key=f"type_{f.name}")
                with c3: 
                    new_year = st.text_input("年度", value=current_config['year'], key=f"year_{f.name}")
                with c4: 
                    new_no = st.selectbox("次別", ["第一次", "第二次", "第三次", "正式考試"], 
                                        index=["第一次", "第二次", "第三次", "正式考試"].index(current_config['exam_no']),
                                        key=f"no_{f.name}")
                
                st.session_state['upload_configs'][f.name] = {
                    "type": new_type, "year": new_year, "exam_no": new_no
                }
                
                final_new_name = f"{new_year}-{new_type}-{new_no}.{f.name.split('.')[-1]}"
                files_to_upload.append({
                    "file_obj": f,
                    "new_filename": final_new_name,
                    "type": new_type,
                    "year": new_year,
                    "exam_no": new_no
                })
            st.divider()

        if st.button("確認並上傳所有檔案", type="primary"):
            duplicate_warnings = []
            for item in files_to_upload:
                existing = cloud_manager.check_file_exists(item['new_filename'])
                if existing:
                    duplicate_warnings.append(f"{item['new_filename']} (原: {item['file_obj'].name})")
            
            if duplicate_warnings:
                st.error(f"發現雲端已有重複檔名，請修改年度或次別：\n" + "\n".join(duplicate_warnings))
            else:
                progress_bar = st.progress(0)
                success_count = 0
                for idx, item in enumerate(files_to_upload):
                    f = item['file_obj']
                    new_fname = item['new_filename']
                    f.seek(0)
                    file_bytes = f.read()
                    
                    # 使用 upload_bytes 回傳的 (url, blob_name)
                    backup_url, blob_name = cloud_manager.upload_bytes(
                        file_bytes, 
                        new_fname, 
                        folder="raw_uploads", 
                        content_type=f.type
                    )
                    
                    file_record = {
                        "filename": new_fname,
                        "original_filename": f.name,
                        "url": backup_url,
                        "blob_name": blob_name, # 儲存 Blob Name
                        "exam_type": item['type'],
                        "year": item['year'],
                        "exam_no": item['exam_no'],
                        "ai_status": "未辨識",
                        "created_at": datetime.datetime.now()
                    }
                    cloud_manager.save_file_record(file_record)
                    
                    st.session_state['file_queue'][new_fname] = {
                        "status": "uploaded", 
                        "data": file_bytes,
                        "type": f.type.split('/')[-1] if '/' in f.type else 'pdf',
                        "result": [],
                        "error_msg": "",
                        "source_tag": f"{item['type']}-{item['year']}",
                        "backup_url": backup_url,
                        "blob_name": blob_name,
                        "db_id": file_record['id'] 
                    }
                    success_count += 1
                    progress_bar.progress((idx + 1) / len(files_to_upload))
                
                if success_count > 0:
                    st.success(f"成功上傳 {success_count} 個檔案！")
                    st.session_state['upload_configs'] = {}
                    time.sleep(1)
                    st.rerun()

    if st.session_state['file_queue']:
        with st.expander(f"查看目前工作階段暫存 ({len(st.session_state['file_queue'])})"):
            for fname in st.session_state['file_queue']:
                st.write(fname)

# === Tab 2: 檔案管理及AI辨識 ===
with tab_files:
    if 'just_processed_file' in st.session_state:
        st.success(f"🎉 **{st.session_state['just_processed_file']}** 辨識完成！")
        st.info("👉 請點選上方 **「📝 AI匯入校對」** 分頁進行檢查。")
        del st.session_state['just_processed_file']

    st.subheader("已上傳考古題檔案庫")
    cloud_files = cloud_manager.load_file_records()
    
    if not cloud_files:
        st.info("目前沒有已上傳的檔案記錄。")
    else:
        files_tree = {}
        for f in cloud_files:
            ftype = f.get('exam_type', '未分類')
            fyear = f.get('year', '未知年份')
            
            if ftype not in files_tree: files_tree[ftype] = {}
            if fyear not in files_tree[ftype]: files_tree[ftype][fyear] = []
            
            files_tree[ftype][fyear].append(f)

        for ftype in sorted(files_tree.keys()):
            with st.expander(f"📁 {ftype}", expanded=False):
                years_dict = files_tree[ftype]
                
                def year_sort_key(y_str):
                    return -int(y_str) if y_str.isdigit() else 0
                
                for fyear in sorted(years_dict.keys(), key=year_sort_key):
                    with st.expander(f"📁 {fyear} 年度", expanded=False):
                        files_list = years_dict[fyear]
                        
                        exam_no_order = {"第一次": 1, "第二次": 2, "第三次": 3, "正式考試": 4, "其他": 99}
                        def file_sort_key(f):
                            no = f.get('exam_no', '其他')
                            return exam_no_order.get(no, 100)
                        
                        sorted_files = sorted(files_list, key=file_sort_key)
                        
                        for f_record in sorted_files:
                            c_name, c_status, c_action = st.columns([5, 2, 3], vertical_alignment="center")
                            
                            with c_name:
                                st.write(f"📄 {f_record.get('filename')}")
                            
                            with c_status:
                                status = f_record.get('ai_status', '未辨識')
                                if status == '已辨識':
                                    st.button("✅ 已辨識", key=f"status_{f_record['id']}", disabled=True, use_container_width=True)
                                else:
                                    st.button("⬜ 未辨識", key=f"status_{f_record['id']}", disabled=True, use_container_width=True)
                            
                            with c_action:
                                b1, b2 = st.columns(2)
                                with b1:
                                    btn_label = "重新辨識" if status == '已辨識' else "AI 辨識"
                                    if st.button(btn_label, key=f"ai_{f_record['id']}", use_container_width=True):
                                        fname = f_record['filename']
                                        
                                        # 嘗試載入檔案
                                        loaded_success = False
                                        blob_name = f_record.get('blob_name')
                                        file_url = f_record.get('url')
                                        
                                        if fname not in st.session_state['file_queue']:
                                            # 優先使用 blob_name 下載 (更穩定)
                                            if blob_name:
                                                file_bytes = cloud_manager.download_blob(blob_name)
                                                if file_bytes:
                                                    st.session_state['file_queue'][fname] = {
                                                        "status": "uploaded", 
                                                        "data": file_bytes,
                                                        "type": fname.split('.')[-1].lower(),
                                                        "result": [],
                                                        "error_msg": "",
                                                        "source_tag": f"{ftype}-{fyear}",
                                                        "backup_url": f_record.get('url'),
                                                        "blob_name": blob_name,
                                                        "db_id": f_record['id']
                                                    }
                                                    loaded_success = True
                                                else:
                                                    st.error("Blob 下載失敗")
                                            # Fallback
                                            elif file_url:
                                                try:
                                                    resp = requests.get(file_url)
                                                    if resp.status_code == 200:
                                                        st.session_state['file_queue'][fname] = {
                                                            "status": "uploaded", 
                                                            "data": resp.content,
                                                            "type": fname.split('.')[-1].lower(),
                                                            "result": [],
                                                            "error_msg": "",
                                                            "source_tag": f"{ftype}-{fyear}",
                                                            "backup_url": f_record.get('url'),
                                                            "db_id": f_record['id']
                                                        }
                                                        loaded_success = True
                                                except: pass
                                        else:
                                            loaded_success = True
                                            
                                        if loaded_success:
                                            process_single_file(fname, api_key_input, f_record['id'])
                                        else:
                                            st.error("無法讀取檔案，請嘗試重新上傳。")

                                with b2:
                                    if st.button("🗑️", key=f"del_f_{f_record['id']}", type="primary", use_container_width=True):
                                        cloud_manager.delete_file_record(f_record['id'])
                                        st.rerun()

# === Tab 3: AI匯入校對 ===
with tab_review:
    st.subheader("匯入校對與截圖")
    ready_files = [f for f, info in st.session_state['file_queue'].items() if info['status'] == 'done']
    
    if not ready_files:
        st.warning("沒有已完成辨識的檔案。請先至「檔案管理及AI辨識」點擊辨識，或上傳新檔案。")
    else:
        default_idx = 0
        if 'just_processed_file' in st.session_state and st.session_state['just_processed_file'] in ready_files:
             default_idx = ready_files.index(st.session_state['just_processed_file'])

        selected_file = st.selectbox("選擇要處理的檔案", ready_files, index=default_idx)
        file_info = st.session_state['file_queue'][selected_file]
        candidates = file_info['result']
        
        st.markdown(f"**正在編輯：{selected_file} (共 {len(candidates)} 題)**")
        
        col_src1, col_src2 = st.columns(2)
        with col_src1:
            default_tag = file_info.get("source_tag", "未分類")
            source_tag = st.text_input("設定此批試卷來源標籤", value=default_tag)
        
        st.divider()
        
        with st.form(key=f"edit_form_{selected_file}"):
            for i, cand in enumerate(candidates):
                st.markdown(f"**第 {cand.number} 題**")
                
                if cand.q_type == "Group":
                    st.info("📖 題組共用敘述")
                
                c1, c2 = st.columns([1, 1])
                with c1:
                    cand.content = st.text_area(f"題目內容 #{i}", cand.content, height=100, key=f"{selected_file}_c_{i}")
                    
                    if cand.q_type != "Group":
                        opts_text = "\n".join(cand.options)
                        new_opts = st.text_area(f"選項 #{i}", opts_text, height=80, key=f"{selected_file}_o_{i}")
                        cand.options = new_opts.split('\n') if new_opts else []
                    
                    # 題型選擇 (中文)
                    current_type_zh = TYPE_MAP_EN_TO_ZH.get(cand.q_type, "單選")
                    new_type_zh = st.selectbox(f"題型 #{i}", TYPE_OPTIONS, index=TYPE_OPTIONS.index(current_type_zh), key=f"{selected_file}_t_{i}")
                    cand.q_type = TYPE_MAP_ZH_TO_EN[new_type_zh]

                    # 題組子題編輯
                    if cand.q_type == "Group" and cand.sub_questions:
                        with st.expander("編輯子題目"):
                            for sub_q in cand.sub_questions:
                                st.text_area(f"子題 {sub_q.get('number')} 內容", sub_q.get('content', ''), key=f"sub_c_{selected_file}_{i}_{sub_q.get('number')}")

                    ans_key = f"{selected_file}_ans_{i}"
                    default_ans = st.session_state.get(ans_key, "")
                    st.text_input(f"答案 (可留空) #{i}", value=default_ans, key=ans_key)
                    
                    chap_idx = 0
                    if cand.predicted_chapter in smart_importer.PHYSICS_CHAPTERS_LIST:
                        chap_idx = smart_importer.PHYSICS_CHAPTERS_LIST.index(cand.predicted_chapter)
                    cand.predicted_chapter = st.selectbox(f"章節分類 #{i}", smart_importer.PHYSICS_CHAPTERS_LIST, index=chap_idx, key=f"{selected_file}_ch_{i}")
                    
                    if cand.image_bytes: st.image(cand.image_bytes, caption="目前附圖", width=200)
                    else: st.caption("🚫 目前無附圖")

                with c2:
                    st.markdown("✂️ **截圖工具**")
                    # 優先使用 ref_image (AI 截取區域)，若無則用 full_page (整頁)
                    image_to_crop = cand.ref_image_bytes if cand.ref_image_bytes else cand.full_page_bytes
                    
                    if image_to_crop:
                        try:
                            # 顯示裁切器
                            if st_cropper:
                                pil_ref = Image.open(io.BytesIO(image_to_crop))
                                st_cropper(
                                    pil_ref, realtime_update=True, box_color='#FF0000',
                                    key=f"{selected_file}_cropper_{i}", aspect_ratio=None
                                )
                                # 這裡僅供互動，實際裁切需要額外邏輯
                                st.caption("可調整紅框範圍")
                            else:
                                st.error("Streamlit Cropper 元件載入失敗，顯示靜態圖片。")
                                st.image(image_to_crop, caption="原始圖片 (無法裁切)")
                        except: 
                            st.error("截圖載入失敗")
                            # 萬一載入失敗，至少顯示靜態圖
                            st.image(image_to_crop, caption="靜態預覽", width=300)
                    else:
                        st.info("無法取得此題的參考圖片 (也無整頁圖片)")
                st.divider()
            
            st.form_submit_button("💾 暫存所有修改 (不會上傳)")
        
        if st.button(f"✅ 確認匯入 [{selected_file}] 至雲端", type="primary"):
            progress_bar = st.progress(0)
            count = 0
            total = len(candidates)
            db_file_id = file_info.get("db_id")

            for i, cand in enumerate(candidates):
                ans_val = st.session_state.get(f"{selected_file}_ans_{i}", "")
                new_q = Question(
                    q_type=cand.q_type,
                    content=cand.content,
                    options=cand.options,
                    source=source_tag, 
                    chapter=cand.predicted_chapter,
                    image_data=cand.image_bytes,
                    answer=ans_val,
                    source_file_id=db_file_id,
                    sub_questions=[Question.from_dict(sq) for sq in cand.sub_questions] if cand.sub_questions else []
                )
                cloud_manager.save_question(new_q.to_dict())
                st.session_state['question_pool'].append(new_q)
                count += 1
                progress_bar.progress((i + 1) / total)
            
            st.success(f"成功匯入 {count} 題！")
            st.session_state['file_queue'][selected_file]['status'] = 'imported'
            if db_file_id:
                cloud_manager.update_file_status(db_file_id, "已匯入")
            st.rerun()

# === Tab 4: 題庫管理與試卷輸出 ===
with tab_bank:
    st.subheader("題庫總覽與試卷輸出")
    if not st.session_state['question_pool']:
        st.info("目前沒有題目。")
    else:
        all_sources = sorted(list(set([q.source for q in st.session_state['question_pool']])))
        selected_questions_for_export = []
        for src in all_sources:
            qs_in_src = [q for q in st.session_state['question_pool'] if q.source == src]
            with st.expander(f"📁 {src} ({len(qs_in_src)} 題)"):
                if st.checkbox(f"選取全套 [{src}] 進行匯出", key=f"sel_src_{src}"):
                    selected_questions_for_export.extend(qs_in_src)
                for i, q in enumerate(qs_in_src):
                    type_badge = TYPE_MAP_EN_TO_ZH.get(q.type, q.type)
                    if q.parent_id: continue 
                    st.markdown(f"**【{type_badge}】 {q.content[:30]}...**")
                    if q.image_url: st.caption("🖼️ 雲端圖片")
                    elif q.image_data: st.caption("💾 本機圖片 (未同步)")
                    with st.popover("編輯"):
                        q.content = st.text_area("題目", q.content, key=f"edt_c_{q.id}")
                        q.answer = st.text_input("答案", q.answer, key=f"edt_a_{q.id}")
                        if st.button("儲存", key=f"save_{q.id}"):
                            cloud_manager.save_question(q.to_dict())
                            st.rerun()
                        if st.button("刪除", key=f"del_{q.id}", type="primary"):
                            cloud_manager.delete_question(q.id)
                            st.rerun()
                    st.divider()

        st.divider()
        st.subheader(f"已選取 {len(selected_questions_for_export)} 題準備匯出")
        if st.button("生成 Word 試卷"):
            f1, f2 = generate_word_files(selected_questions_for_export)
            st.download_button("下載試題卷", f1, "exam.docx")
            st.download_button("下載答案卷", f2, "ans.docx")
